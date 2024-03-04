"""
This code is used to generate the reports using the normalized json survey 
files.
"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx import oxml
from docx.shared import Pt, Cm, RGBColor
import pandas as pd
import os

from graph_generation import getGraphs
from read_json import getSurveyData
from json_cleaning import json_cleaning


def addImageToCell(cell, imagePath):
    """
    Adds an image to a table cell.

    Parameters
    ----------
    cell : table cell
        The cell in which the iamge will be added.
    imagePath : string
        The path to the image.
    """
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    format = paragraph.paragraph_format
    format.space_before = Pt(0)
    format.space_after = Pt(0)
    paragraph.add_run().add_picture(imagePath, width=Cm(8))
    

def getCompleteAnswerRate(questions, nbStudents):
    """
    Computes the complete answer rate of the questions. The complete answer 
    rate is the portion of students who answered all questions in percent.

    Parameters
    ----------
    questions : dict
        A dictionnary containing the list of questions.
    nbStudents : int
        The number of students in the course.
    
    Returns
    -------
    out : int
        The complete answer rate.
    """
    if len(questions) == 0:
        return 0
    if nbStudents == 0:
        return 100
    return round(sum(min(questions.values(), key=lambda x:sum(x['answer'][0:5]))['answer'][0:5]) / nbStudents * 100)


def getPartialAnswerRate(questions, nbStudents):
    """
    Computes the partial answer rate of the questions. The partiel answer 
    rate is the portion of students who answered at least one question in 
    percent.

    Parameters
    ----------
    questions : dict
        A dictionnary containing the list of questions.
    nbStudents : int
        The number of students in the course.
    
    Returns
    -------
    out : int
        The partial answer rate.
    """
    if len(questions) == 0:
        return 0
    if nbStudents == 0:
        return 100
    return round(sum(max(questions.values(), key=lambda x:sum(x['answer'][0:5]))['answer'][0:5]) / nbStudents * 100)


def getCourseInfo(courseCode, courseTable):
    """
    Gets the formation, the year, the cursus and the acronym of a given course. 

    Parameters
    ----------
    courseCode : string
        The code identifying the course.
    courseTable : Dataframe
        The course table.
    
    Returns
    -------
    out : tuple of 4 strings
        The formation, the year, the cursus and the acronym of the course.
    """
    return courseTable.loc[courseTable['surveyCode'] == courseCode][['year', 'cursus', 'acronym']].iloc[0]


def getNbStudents(courseCode, courseTable):
    """
    Gets the number of students in a given course. 

    Parameters
    ----------
    courseCode : string
        The code identifying the course.
    courseTable : Dataframe
        The course table.
    
    Returns
    -------
    out : int
        The number of students in the course. 
    """
    return courseTable.loc[courseTable['surveyCode'] == courseCode]['nbStudents'].iloc[0]


def getSectionName(sectionCode, sectionTable):
    """
    Gets the section name associated with a section code. If the section code 
    does not match any section, returns "Autres questions".

    Parameters
    ----------
    sectionCode : string
        The code identifying the section.
    sectionTable : Dataframe
        The section table.
    
    Returns
    -------
    out : string
        The name of the section
    """
    if not(sectionCode in sectionTable['code'].unique()):
        return "Autres questions"
    return sectionTable.loc[sectionTable['code'] == sectionCode]['name'].iloc[0]

def verifyCourse(courseCode, courseTable):
    try:
        ans = courseTable.loc[courseTable['surveyCode'] == courseCode]['nbStudents'].iloc[0]
        return True
    except:
        return False

def addPositiveRateInfo(question, nbAnswers):
    """
    Adds the positive rate info to a question. The positive rate is considered 
    zero if the number of answers is zero.

    Parameters
    ----------
    question : dict
        A dictionnary representing the question.
    nbAnswers : int
        The number of answers to the question.
    """
    if nbAnswers != 0:
        question['positiveRate'] = round(sum(question['answer'][2:4]) / nbAnswers * 100)
    else:
        question['positiveRate'] = 0


def addAnswerRateInfo(question, nbAnswers, nbStudents):
    """
    Adds the answers rate info to a question. The answers rate is considered 
    zero if the number of answers is zero.

    Parameters
    ----------
    question : dict
        A dictionnary representing the question.
    nbAnswers : int
        The number of answers to the question.
    nbStudents : int
        The number of students in the course.
    """
    if nbStudents != 0:
        question['answerRate'] = round(nbAnswers / nbStudents * 100)
    else:
        question['answerRate'] = 0


def addRatesInfo(questions, nbStudents):
    """
    Adds the answers rate and positive rate info to a question.

    Parameters
    ----------
    question : dict
        A dictionnary representing the question.
    nbStudents : int
        The number of students in the course.
    """
    for _, question in questions.items():
        if question['isClosed']:
            answer = question['answer']
            nbAnswers = sum(answer[0:4])
            addPositiveRateInfo(question, nbAnswers)
            addAnswerRateInfo(question, nbAnswers, nbStudents)


def getGoodBadPositiveRates(questions, minAnswerRate, minGoodPositiveRate, maxBadPositiveRate, nbPoints=3):
    """
    Gets the values of the good and bad positive rates in order to have a 
    specific number of good questions and bad questions

    Parameters
    ----------
    questions : dict
        A dictionnary representing the question.
    minAnswerRate : int
        The minimum answer rate for a question to be considered in the 
        computation.
    minGoodPositiveRate : int
        The minimum positive rate for a question to be considered good.
    maxBadPositiveRate : int
        The maximum positive rate for a question to be considered bad.
    nbPoints : int, optional
        The number of points.
    
    Returns
    -------
    out : tuple of 2 int
        The good and bad positive rates.
    """
    consideredQuestions = {key: question for key, question in questions.items() if question['isClosed'] and question['answerRate'] >= minAnswerRate}
    sortedQuestions = sortDictListByKey(consideredQuestions.values(), 'positiveRate')
    if len(sortedQuestions) == 0:
        return minGoodPositiveRate, maxBadPositiveRate
    if len(sortedQuestions) <= nbPoints:
        return sortedQuestions[0]['positiveRate'], sortedQuestions[-1]['positiveRate']
    return sortedQuestions[-nbPoints]['positiveRate'], sortedQuestions[nbPoints]['positiveRate']
    

def getAveragePositiveRate(questions, minAnswerRate):
    """
    Computes the average positive rate on a list of questions for questions 
    with an answer rate greater than minAnswerRate.

    Parameters
    ----------
    questions : dict
        A dictionnary of questions.
    minAnswerRate : int
        The minimum answer rate for a question to be considered in the 
        computation.
    
    Returns
    -------
    out : int
        The average positive rate.
    """
    totalPositiveRate = 0
    nbQuestions = 0
    for _, question in questions.items():
        if question['answerRate'] >= minAnswerRate:
            totalPositiveRate += question['positiveRate']
            nbQuestions += 1
    if nbQuestions != 0:
        return round(totalPositiveRate / nbQuestions)
    else:
        return 0


def getGlobalQualityPositiveRate(questions):
    """
    Gets the positive rate of the the global quality question.
    If the global quality question is not found, returns 0.

    Parameters
    ----------
    questions : dict
        A dictionnary of questions.
    
    Returns
    -------
    out : int
        The global positive rate or 0 if the global quality question is not 
        found.
    """
    minID = -1
    positiveRate = 0
    for key, question in questions.items():
        if question['section'] == "AGXX":
            id = int(key[-3:])
            if minID == -1 or id < minID:
                positiveRate = question['positiveRate']
                minID = id
    return positiveRate


def sortDictListByKey(dictList, key):
    """
    Sorts a list of dictionnaries on a specific key.

    Parameters
    ----------
    dictList : list of dict
        A list of dictionnaries.
    key : string
        The key.
    
    Returns
    -------
    out : int
        The sorted list.
    """
    return sorted(dictList, key=lambda q: q[key])


def getGoodBadPoints(questions, minAnswerRate, goodPositiveRate, badPositiveRate):
    """
    Gets the list of good and bad points. A question is considered bad if its 
    positive rate is below maxBadPositiveRate. A question is considered 
    good if its positive rate is greater than minGoodPositiveRate.

    Parameters
    ----------
    questions : dict
        A dict of questions.
    minAnswerRate : int
        The minimum answer rate for a question to be considered in the 
        computation.
    goodPositiveRate : int
        The positive rate for a question to be considered good.
    badPositiveRate : int
        The positive rate for a question to be considered bad.
    
    Returns
    -------
    out : tuple of 2 lists
        A tuple containing the list of good points and the list of bad points.
    """
    goodPoints, badPoints = [], []
    for _, question in questions.items():
        if question['answerRate'] >= minAnswerRate:
            if question['positiveRate'] >= goodPositiveRate:
                goodPoints.append(question)
            elif question['positiveRate'] <= badPositiveRate:
                badPoints.append(question)
    goodPoints = sortDictListByKey(goodPoints, 'positiveRate')
    badPoints = sortDictListByKey(badPoints, 'positiveRate')
    goodPoints.reverse()
    return goodPoints, badPoints


def setLanguageToFrench(document):
    """
    Sets the language of a document to french.

    Parameters
    ----------
    document : Document
        The document.
    """
    for my_style in document.styles:
        style = document.styles[my_style.name]
        rpr = style.element.get_or_add_rPr()
        lang = oxml.shared.OxmlElement('w:lang')
        if not rpr.xpath('w:lang'):
            lang.set(oxml.shared.qn('w:val'),'fr-FR')
            lang.set(oxml.shared.qn('w:eastAsia'),'en-US')
            lang.set(oxml.shared.qn('w:bidi'),'ar-SA')
            rpr.append(lang)


def setMargin(document, margin):
    """
    Sets the margin of the document.

    Parameters
    ----------
    document : Document
        The document.
    margin : Cm
        The margin in centimeters.
    """
    for section in document.sections:
        section.left_margin = margin
        section.right_margin = margin


def setStyles(document):
    """
    Sets the styles for the document.

    Parameters
    ----------
    document : Document
        The document.
    """
    # Sets the style for Heading 1
    style = document.styles['Heading 1']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(14)
    font.underline = True
    font.bold = True
    font.color.rgb = RGBColor(0, 0, 0)

    # Sets the style for Heading 2
    style = document.styles['Heading 2']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(13)
    font.bold = True
    font.italic = True
    font.color.rgb = RGBColor(0, 0, 0)

    # Sets the style for Heading 2
    style = document.styles['Heading 3']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)
    font.bold = True
    font.color.rgb = RGBColor(0, 0, 0)

    # Sets the style for Normal
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)


def getFormationName(formationCode):
    """
    Gets the formation name associated to a formation code.

    Paramaters
    ----------
    formationCode : string
        The code of the formation.
    
    Returns
    -------
    out : string
        The formation name.
    """
    if formationCode == "INGE":
        return "INGE GEN"
    return ""


def getCursusName(cursusCode):
    """
    Gets the cursus name associated to a cursus code.

    Paramaters
    ----------
    cursusCode : string
        The code of the cursus.
    
    Returns
    -------
    out : string
        The cursus name.
    """
    if cursusCode == "OD":
        return "Option Disciplinaire"
    if cursusCode == "TC":
        return "Tronc Commun"
    if cursusCode == "SoftSkills":
        return "Soft Skills"
    if cursusCode == "LVC":
        return "LVC"
    if cursusCode == "LVO":
        return "LVO"
    if cursusCode == "EPS":
        return "EPS"
    if cursusCode == "FastTrack":
        return "Fast Track"
    if cursusCode == "PROAR":
        return "Projet Architectural"
    return ""

def getCursusCode(cursus):
    if "BTP" in cursus:
        return "BTP"
    if "Option disciplinaire" in cursus:
        return "OD"
    if "Tronc commun" in cursus:
        return "TC"
    if "Option professionnelle" in cursus:
        return "OP"
    if "MECA" in cursus:
        return "MEC"
    if "SEC" in cursus:
        return "SEC"
    if "Fast Track" in cursus:
        return "FastTrack"
    return "X"

def displayCourseInfo(document, acronym, formation, year, cursus):
    """
    Displays information on the course in the document.

    Paramaters
    ----------
    document : Document
        The document.
    acronym : string
        The acronym of the course.
    formation : string
        The formation of the course.
    year : string
        The year of the course.
    cursus : string
        The cursus of the course.
    """
    text = f"Titre de la matière : {acronym}\n"
    #text += f"Cursus : {getFormationName(formation)}\n"
    text += f"Cursus : Année {year[:2]}/{year[2:]} - {cursus}"
    document.add_paragraph(text)


def generateResponseRateSection(closedQuestions, document, nbStudents):
    """
    Generates the response rate section of the report.

    Parameters
    ----------
    closedQuestions : dict
        A dictionnary containing the closed questions.
    document : Document
        The document.
    nbStudents : int
        The number of students in the course.
    
    Returns
    -------
    out : tuple of 2 int
        A tuple containing the complete answer rate and the partial answer rate.
    """
    document.add_heading("Taux de réponses", level=2)
    detail = document.add_paragraph()
    run = detail.add_run("(moyenne, sur l'ensemble des questions, du nombre de réponses exprimées / nombre d'inscrits dans la matière)")
    run.font.size = Pt(10)
    detail.paragraph_format.space_after = 0
    paragraph = document.add_paragraph(text=f"Le taux de réponse à l'enquête est de (pour un total de {nbStudents} étudiants inscrits) :")
    paragraph.paragraph_format.space_after = 0
    completeAnswerRate = getCompleteAnswerRate(closedQuestions, nbStudents)
    partialAnswerRate = getPartialAnswerRate(closedQuestions, nbStudents)
    document.add_paragraph(f"{completeAnswerRate} % pour les réponses complètes", style='List Bullet')
    document.add_paragraph(f"{partialAnswerRate} % pour les réponses partielles.", style='List Bullet')
    return completeAnswerRate, partialAnswerRate


def generatePositiveRatesSections(document, questions, minAnswerRate, goodPositiveRate, badPositiveRate, maxNbGoodBadPoints):
    """
    Generates the positive rate section of the report.

    Parameters
    ----------
    document : Document
        The document.
    closedQuestions : dict
        A dictionnary containing the closed questions.
    minAnswerRate : int
        The minimum answer rate for a question to be considered in the 
        computation.
    goodPositiveRate : int
        The positive rate for a question to be considered good.
    badPositiveRate : int
        The positive rate for a question to be considered bad.
    maxNbGoodBadPoints : int
        The maximum number of good and bad points

    Returns
    -------
    out : tuple of 2 int
        A tuple containing the number of good points and the number of bad points.
    """
    # Create the global evaluation section
    document.add_heading("Appréciation générale", level=2)
    document.add_paragraph(f"Qualité globale de l'enseignement : {getGlobalQualityPositiveRate(questions)} % des étudiants sont satisfaits.")

    # Create the list of good points and bad points
    goodPoints, badPoints = getGoodBadPoints(questions, minAnswerRate, goodPositiveRate, badPositiveRate)
    
    # Display the lists
    if len(goodPoints) != 0:
        document.add_heading("Les points positifs", level=2)
        detail = document.add_paragraph()
        run = detail.add_run("(% : réponses tout à fait d'accord et plutôt d'accord sur les réponses exprimées)")
        run.font.size = Pt(10)
        detail.paragraph_format.space_after = 0
        for k in range(min(len(goodPoints), maxNbGoodBadPoints)):
            goodPoint = goodPoints[k]
            document.add_paragraph(f"{goodPoint['question']} ({goodPoint['positiveRate']} %)", style='List Bullet')
    if len(badPoints) != 0:
        document.add_heading(f"Les points sensibles", level=2)
        detail = document.add_paragraph()
        run = detail.add_run("(% : réponses pas du tout d'accord et plutôt pas d'accord sur les réponses exprimées)")
        run.font.size = Pt(10)
        detail.paragraph_format.space_after = 0
        for k in range(min(len(badPoints), maxNbGoodBadPoints)):
            badPoint = badPoints[k]
            document.add_paragraph(f"{badPoint['question']} ({100 - badPoint['positiveRate']} %)", style='List Bullet')
    
    return len(goodPoints), len(badPoints)


def getNbSections(questions):
    """
    Gets the number of sections in the questions.

    Parameters
    ----------
    questions : dict
        A dictionnary of questions.

    Returns
    -------
    out : int
        The number of sections.
    """
    sections = set()
    for question in questions:
        if not(question['section'] in sections):
            sections.add(question['section'])
    return len(sections)


def shadeTableCell(cell, hexColor):
    """
    Fills the cell with the provided color.

    Parameters
    ----------
    cell : table cell
        The cell to be shaded.
    hexColor : string
        The hex code for the shading color.
    """
    shading_elm_1 = oxml.parse_xml(r'<w:shd {} w:fill="{}"/>'.format(oxml.ns.nsdecls('w'), hexColor))
    cell._tc.get_or_add_tcPr().append(shading_elm_1)


def addLegend(table, rowIndex):
    """
    Adds a legend to a table.

    Parameters
    ----------
    table : Table
        The table in which the legend is added.
    rowIndex : int
        The index of the row of the legend.
    """
    table.cell(rowIndex, 0).merge(table.cell(rowIndex, 2))
    paragraph = table.cell(rowIndex, 0).paragraphs[0]
    format = paragraph.paragraph_format
    format.space_after = Cm(0.2)
    format.space_before = Cm(0.2)
    run = paragraph.add_run("Légende : ")
    run.bold = True
    colors = {
        " Pas du tout d'accord ": RGBColor(255, 0, 0), 
        " Plutôt pas d'accord ": RGBColor(255, 192, 0), 
        " Plutôt d'accord ": RGBColor(0, 176, 240), 
        " Tout à fait d'accord": RGBColor(146, 208, 80)
    }
    for key, color in colors.items():
        run = paragraph.add_run("\u25A0")
        run.font.color.rgb = color
        paragraph.add_run(key)


def addSectionHeader(table, rowIndex, sectionCode, first, sectionTable):
    """
    Adds the header of a section to a table.

    Parameters
    ----------
    table : Table
        The table in which the header is added.
    rowIndex : int
        The index of the row of the header.
    sectionCode : string
        The code of the section.
    first : boolean
        If True, the header is the first header in the document. Thus, it 
        contains the titles of all the columns.
    sectionTable : Dataframe
        The section table.
    """
    sectionName = getSectionName(sectionCode, sectionTable)
    table.cell(rowIndex, 0).text = sectionName
    table.cell(rowIndex, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    format = table.cell(rowIndex, 0).paragraphs[0].paragraph_format
    format.space_after = Cm(0.2)
    format.space_before = Cm(0.2)
    shadeTableCell(table.cell(rowIndex, 0), 'DDDDDD')
    if first:
        table.cell(rowIndex, 1).text = "Répartition des réponses"
        table.cell(rowIndex, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.cell(rowIndex, 2).text = "Taux de réponse"
        table.cell(rowIndex, 2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def generateQuestionsTable(document, questions, graphs, sectionTable):
    """
    Generates the question table.

    Parameters
    ----------
    document : Document.
        The document.
    questions : dict
        A dictionnary of questions.
    graphs : list
        A list containing the paths to the graph images.
    sectionTable : Dataframe
        The section table.
    """
    # Create the explanation paragraph
    document.add_paragraph(f"Le % affiché sur les couleurs est calculé sur les réponses exprimées.\nLe nombre à droite correspond au % de réponses exprimées (= 100 - % de (non réponses et de sans avis/ne s'applique pas)).")

    # Sort the questions by sections
    for key, question in questions.items():
        question['key'] = key
    questions = sortDictListByKey(questions.values(), 'section')
    nbSections = getNbSections(questions)
    
    # Initialize the table
    table = document.add_table(rows=len(questions) + nbSections + 1, cols=3, style='Table Grid')
    for cell in table.columns[0].cells:
        cell.width = Cm(7)
    for cell in table.columns[1].cells:
        cell.width = Cm(8)
    for cell in table.columns[2].cells:
        cell.width = Cm(2)

    # Create the legend
    addLegend(table, 0)

    # Create the rest of the table
    currentSection = ""
    rowIndex = 1
    for question in questions:
        
        # Add a section header if necessary
        if question['section'] != currentSection:
            currentSection = question['section']
            addSectionHeader(table, rowIndex, currentSection, rowIndex == 1, sectionTable)
            rowIndex += 1
        
        # Add the question text
        table.cell(rowIndex, 0).text = question['question']
        
        # Add the graph
        addImageToCell(table.cell(rowIndex, 1), graphs[question['key']])

        # Add the answer rate
        table.cell(rowIndex, 2).text = f"{question['answerRate']} %"
        paragraph = table.cell(rowIndex, 2).paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(rowIndex, 2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        rowIndex += 1


def generateOpenQuestionsSection(document, questions):
    """
    Generates the open question section of the report.

    Parameters
    ----------
    document : Document.
        The document.
    questions : dict
        A dictionnary of questions.
    """
    # Create the heading
    document.add_heading(f"Liste des commentaires", level=2)

    # Create the answers lists
    for _, question in questions.items():
        heading = document.add_heading(question['question'], level=3)
        format = heading.paragraph_format
        format.space_after = Pt(0)
        for answer in question['answer']:
            document.add_paragraph(answer, style='List Bullet')


def generateReport(courseCode, closedQuestions, openQuestions, graphs, outputFolder, nbStudents, minAnswerRate, goodPositiveRate, badPositiveRate, maxNbGoodBadPoints, metaData, courseTable, sectionTable):
    """
    Generates the report.

    Parameters
    ----------
    courseCode : string.
        The code of the course.
    closedQuestions : dict
        A dictionnary of closed questions.
    openQuestions : dict
        A dictionnary of open questions.
    graphs : list
        A list containing the paths to the graph images.
    outputFolder : string
        The path to the folder in which the report must be saved.
    nbStudents : int
        The number of students in the course.
    minAnswerRate : int
        The minimum answer rate for a question to be considered in the 
        computation.
    goodPositiveRate : int
        The positive rate for a question to be considered good.
    badPositiveRate : int
        The positive rate for a question to be considered bad.
    maxNbGoodBadPoints : int
        The maximum number of good and bad points.
    metaData : dataframe
        Dataframe with some stats on the reports.
    courseTable : Dataframe
        The course table.
    sectionTable : Dataframe
        The section table.
    """
    # Get the information on the course
    year, cursus, acronym = getCourseInfo(courseCode, courseTable)

    # Create the document
    document = Document()
    setLanguageToFrench(document)
    setMargin(document, Cm(2))
    setStyles(document)
    
    # Create the heading
    displayCourseInfo(document, acronym, "INGE", str(year), cursus)
    document.add_heading("1. Synthèse de l'évaluation", level=1)

    # Create the response rate section
    completeAnswerRate, partialAnswerRate = generateResponseRateSection(closedQuestions, document, nbStudents)

    # Create the evaluation sections
    nbGoodPoints, nbBadPoints = generatePositiveRatesSections(document, closedQuestions, minAnswerRate, goodPositiveRate, badPositiveRate, maxNbGoodBadPoints)

    # Add additionnal headings for final report
    document.add_heading("Proposition d'amélioration", level=2)
    document.add_heading("2. Réponse de l'enseignant·e responsable de la matière", level=1)
    document.add_heading("Plan d'action, pilote et échéancier", level=2)

    # Create the table with the questions
    document.add_page_break()
    generateQuestionsTable(document, closedQuestions, graphs, sectionTable)

    # Create the open questions section
    document.add_page_break()
    generateOpenQuestionsSection(document, openQuestions)

    # Save the document in the output folder
    document.save(f"{outputFolder}/INGE_{year}_{getCursusCode(cursus)}_{acronym}_V0.docx")

    # Save the metadata
    metaData.loc[len(metaData.index)] = [acronym, completeAnswerRate, partialAnswerRate, nbGoodPoints, nbBadPoints, getGlobalQualityPositiveRate(closedQuestions)]


def generateReports(jsonFolderPath, outputFolder, minAnswerRate, minGoodPositiveRate, maxBadPositiveRate, isNormalized, maxNbGoodBadPoints=4):
    """
    Generates all the report for a list of JSON files.

    Parameters
    ----------
    jsonFolderPath : string.
        The path to the folder of JSON files.
    outputFolder : string
        The path to the folder in which the report must be saved.
    minAnswerRate : int
        The minimum answer rate for a question to be considered in the 
        computation.
    minGoodPositiveRate : int
        The minimum positive rate for a question to be considered good.
    maxBadPositiveRate : int
        The maximum positive rate for a question to be considered bad.
    isNormalized : boolean
        If True, the json files in the json folder are normalized.
    maxNbGoodBadPoints : int, optional
        The maximum number of good and bad points. Defaults to 4.
    """
    # Normalize the data if necessary
    if not(isNormalized):
        if os.path.exists("temp_normalizedJSON"):
            for path in os.listdir("temp_normalizedJSON"):
                os.remove(os.path.join("temp_normalizedJSON", path))
        else:
            os.makedirs("temp_normalizedJSON")
        json_cleaning(jsonFolderPath, "temp_normalizedJSON")
        jsonFolderPath = "temp_normalizedJSON"

    # Recover the data and initialize the metadata
    surveyData = getSurveyData(jsonFolderPath)
    courseTable = pd.read_csv("course_table.csv")
    sectionTable = pd.read_csv("section_table.csv")
    metaData = pd.DataFrame(columns = ['courseName', 'completeAnswerRate', 'partialAnswerRate', 'nbGoodPoints', 'nbBadPoints', 'globalQuality'])

    # Get the good and bad positive rates
    averageGoodPositiveRate, averageBadPositiveRate = 0, 0
    lenSurveyData = 0
    for courseCode in surveyData:
        if (verifyCourse(courseCode, courseTable)):
            nbStudents = getNbStudents(courseCode, courseTable)
            addRatesInfo(surveyData[courseCode], nbStudents)
            goodPositiveRate, badPositiveRate = getGoodBadPositiveRates(surveyData[courseCode], minAnswerRate, minGoodPositiveRate, maxBadPositiveRate)
            averageGoodPositiveRate += goodPositiveRate
            averageBadPositiveRate += badPositiveRate
            lenSurveyData += 1
    goodPositiveRate = max(round(averageGoodPositiveRate / max(lenSurveyData,1)), minGoodPositiveRate)
    badPositiveRate = min(round(averageBadPositiveRate / max(lenSurveyData,1)), maxBadPositiveRate)

    # Generate the reports
    count = 1
    for courseCode in surveyData:
        if (verifyCourse(courseCode, courseTable)):
            print(count, courseCode, "yes")
            closedQuestions = {key: question for key, question in surveyData[courseCode].items() if question['isClosed']}
            openQuestions = {key: question for key, question in surveyData[courseCode].items() if not(question['isClosed'])}
            
            graphs = getGraphs(closedQuestions)
            nbStudents = getNbStudents(courseCode, courseTable)
            generateReport(courseCode, closedQuestions, openQuestions, graphs, outputFolder, nbStudents, minAnswerRate, goodPositiveRate, badPositiveRate, maxNbGoodBadPoints, metaData, courseTable, sectionTable)
        else:
            print(count, courseCode, "no")
        count += 1
        # if count > 5:
        #     break
    # Complete and save the metadata
    metaData.loc[len(metaData.index)] = ['Average'] + list(metaData.mean(numeric_only=True))
    metaData.to_csv(f"{outputFolder}/metaAnalysis.csv", index=False)

    # Remove the temporary normalized json files
    if not(isNormalized):
        for file in os.listdir(jsonFolderPath):
            os.remove(os.path.join(jsonFolderPath, file))
        os.rmdir(jsonFolderPath)